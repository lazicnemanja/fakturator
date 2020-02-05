#!/usr/bin/env python3

import sys
import os
import time

from docx import Document
from nbs import NBS
from fakturator import Fakturator

nbs_url = "https://www.nbs.rs/kursnaListaModul/srednjiKurs.faces"

def cls():
    os.system('cls' if os.name=='nt' else 'clear')

def get_argv():
    if len(sys.argv) > 1: 
        try:
            return float(sys.argv[1])
        except ValueError:
            print ("Wrong type, '{}' must be a number!").format(sys.argv[1])
    else:
        print ("Missing param!, enter a value in EUR!")
        exit()

def need_to_edit(placeholders):
    cls()
    print ('-'*30)
    for k in sorted(placeholders.keys()):
        print ("{} = {}".format(k, placeholders[k]))

    proceed = False;
    while True:
        i = input("\nLooking good? (Y/n)").lower()
        if i == '' or i == 'y':
            proceed = True
            break
        elif i == 'n':
            proceed = False
            break

    return proceed

def edit(placeholders):
    for k in sorted(placeholders.keys()):
        cls()
        i=input("{}({}) = ".format(k, placeholders[k])).lower()
        if i != '':
            placeholders[k] = i;
        else:
            pass

def main():
    doc = Document('./Template.docx')
    nbs = NBS(nbs_url)
    fakt = Fakturator()
    user_eur = get_argv()
    
    placeholders = {
        "{invoice-no}": fakt.get_invoice_no(),
        "{completion-date}": fakt.get_today(),
        "{issued-date}": fakt.get_today(),
        "{deadline-date}": fakt.get_deadline_date(),
        "{from-date}": fakt.get_date_from(),
        "{to-date}":fakt.get_date_to(),
        "{nbs}" : str(nbs.eur_exchange_rate()),
        "{nbs-date}": nbs.formed_on(),
        "{eur}" : str(round(user_eur,2)), 
        "{rsd}" : str(round(user_eur * nbs.eur_exchange_rate(),2))
    }

    proceed = False
    while proceed == False:
        proceed = need_to_edit(placeholders)
        if proceed == False:
            edit(placeholders)
        else:
            break

    for p in doc.paragraphs:
        for k in placeholders.keys():
            if k in p.text:
                p.text = p.text.replace(k,placeholders[k])


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k in placeholders.keys():
                        if k in p.text:
                            p.text = p.text.replace(k,placeholders[k])

    doc.save("./{}.docx".format(placeholders['{invoice-no}']))
    print ("Done!")


if __name__ == "__main__":
    main()